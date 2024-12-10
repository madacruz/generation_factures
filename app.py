import streamlit as st
import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import re
from io import BytesIO
from datetime import datetime
from docx.shared import Pt
import zipfile
import subprocess

# https://github.com/Franky1/Streamlit-docx-converter

# Configuration de la page Streamlit
st.set_page_config(page_title="Générateur de Factures Grands Formats", layout="wide")

# Création des dossiers pour les fichiers
os.makedirs("factures_docx", exist_ok=True)
os.makedirs("factures_pdf", exist_ok=True)

def replace_text_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value)
            if key == "{{NUMERO}}":
                for run in paragraph.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True

def safe_filename(filename):
    return re.sub(r'[<>:"/\\\\|?*]', '-', filename)

def convert_to_pdf_with_libreoffice(docx_path, pdf_dir):
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to", "pdf:writer_pdf_Export",
                "--outdir", pdf_dir,
                docx_path
            ],
            check=True
        )
        pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        return os.path.join(pdf_dir, pdf_filename)
    except subprocess.CalledProcessError as e:
        st.error(f"Erreur lors de la conversion avec LibreOffice : {e}")
        return None

def generer_facture(row, template_path, numero_facture, date_facture):
    doc = Document(template_path)
    replacements = {
        "{{NOM}}": str(row['NOM']),
        "{{PRENOM}}": str(row['PRENOM']),
        "{{STRUCTURE}}": str(row['STRUCTURE']),
        "{{ENSEMBLE}}": str(row['ENSEMBLE']),
        "{{TARIF}}": f"{int(row['TARIF'])}",
        "{{NUMERO}}": str(numero_facture),
        "{{DATE}}": date_facture.strftime("%d/%m/%Y")
    }

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for table_row in table.rows: 
            for cell in table_row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    docx_filename = f"2024-ADHF{numero_facture} - {row['ENSEMBLE']} cotisation annuelle.docx"
    docx_filename = safe_filename(docx_filename)
    docx_path = os.path.join("factures_docx", docx_filename)
    doc.save(docx_path)

    pdf_filename = f"2024-ADHF{numero_facture} - {row['ENSEMBLE']} cotisation annuelle.pdf"
    pdf_filename = safe_filename(pdf_filename)
    pdf_path = os.path.join("factures_pdf", pdf_filename)

    try:
        convert(docx_path, pdf_path)
    except Exception:
        try:
            pdf_dir = "factures_pdf"
            os.makedirs(pdf_dir, exist_ok=True)
            pdf_path = convert_to_pdf_with_libreoffice(docx_path, pdf_dir)
        except Exception as e:
            pdf_path = None  # Ne pas interrompre le processus si la conversion échoue
    return docx_path, pdf_path

def capitalize_name(name):
    parts = [part.capitalize() for part in name.split(' ')]
    capitalized_name = ' '.join(['-'.join([subpart.capitalize() for subpart in part.split('-')]) for part in parts])
    return capitalized_name

st.title("Générateur de Factures Grands Formats")

col1, col2 = st.columns(2)

with col1:
    uploaded_file = st.file_uploader("Téléchargez votre fichier CSV", type="csv")
    indice_depart = st.number_input("Indice de départ pour les factures", min_value=1, value=1, step=1)
    date_facture = st.date_input("Date de la facture", value=datetime.today())

if uploaded_file:
    df_original = pd.read_csv(uploaded_file)
    
    df = df_original.rename({
        "Nom de la structure juridique": "STRUCTURE",
        "Nom du ou des ensemble(s) et/ou collectif membre(s) de Grands Formats": "ENSEMBLE",
        "Nom du référent": "NOM",
        "Prénom du référent": "PRENOM",
        "Le montant de ma cotisation est de :\nPour un budget :\n- inférieur à 10 000 euros : 75 euros\n- compris entre 10 000 et 85 000 euros : 150 euros\n- compris entre 85 000 et 150 000 euros : 250 euros\n- supérieur à 150 000 euros : 350 euros)": "TARIF"
    }, axis=1)
    
    df = df[["STRUCTURE", "ENSEMBLE", "NOM", "PRENOM", "TARIF"]]
    df['NOM'] = df['NOM'].apply(capitalize_name)
    df['PRENOM'] = df['PRENOM'].apply(capitalize_name)
    df['TARIF'] = df['TARIF'].apply(lambda x: int(re.search(r'\d+', str(x)).group()) if pd.notnull(x) else 0)

    with col2:
        if st.button("Générer les factures"):
            pdf_files = []
            docx_files = []
    
            progress = st.progress(0)  # Initialisation de la barre de progression
            total_rows = len(df)
            
            for i, row in enumerate(df.iterrows(), 0):
                numero_facture = indice_depart + i
                docx_path, pdf_path = generer_facture(row[1], 'Modèle facture cotisations.docx', numero_facture, date_facture)
                docx_files.append(docx_path)
                if pdf_path:
                    pdf_files.append(pdf_path)
                
                progress.progress(i / total_rows)
            
            st.success(f"Factures générées avec succès ! DOCX : {len(docx_files)}, PDF : {len(pdf_files)}.")
            
            zip_buffer = BytesIO()
            with st.spinner("Compression des fichiers..."):
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    for file in docx_files + pdf_files:
                        zf.write(file, os.path.relpath(file))
            
            st.download_button(
                label="Télécharger toutes les factures",
                data=zip_buffer.getvalue(),
                file_name="factures.zip",
                mime="application/zip"
            )

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Aperçu avant modifications")
        st.write(df_original.head(50))
    with col2:
        st.subheader("Aperçu après modifications")
        st.write(df.head(50))
